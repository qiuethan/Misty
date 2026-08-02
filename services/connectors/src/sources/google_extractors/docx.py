"""Uploaded-.docx extraction to markdown.

Downloads bytes through the same Drive media path as text/* uploads and parses
locally, so no Google scope beyond Drive and no new API client. Output mirrors
the Docs extractor, because a .docx is the same kind of document as a Google
Doc.

Legacy .doc (application/msword) is a different binary format python-docx
cannot read; it stays SourceUnsupported.
"""

import io

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.sources.base import SourceUnsupported
from src.sources.google_extractors.base import ExtractedText, execute
from src.sources.google_extractors.drive_export import DRIVE_READONLY
from src.sources.google_extractors.markdown import render_markdown_table

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

_HEADING_PREFIX = {
    "Title": "#",
    "Heading 1": "#",
    "Heading 2": "##",
    "Heading 3": "###",
    "Heading 4": "####",
    "Heading 5": "#####",
    "Heading 6": "######",
}


def _paragraph_line(paragraph) -> str:
    text = (paragraph.text or "").strip()
    if not text:
        return ""
    style = getattr(paragraph.style, "name", "") or ""
    prefix = _HEADING_PREFIX.get(style)
    if prefix:
        return f"{prefix} {text}"
    if _is_bullet_style(style):
        return f"- {text}"
    return text


def _is_bullet_style(style: str) -> bool:
    """True only for styles that render as an actual bulleted/numbered item.

    A raw "List" in style substring check false-positives on styles that
    contain the word but are not bullets: "List Continue"/"List Continue
    2/3" are prose continuing under a list item, deliberately without its
    own bullet, and plain "List" is just an indented paragraph. "List
    Paragraph" is what Word applies when you click the bullet button and is
    the most common style in practice, so it is matched exactly rather than
    by prefix.
    """
    return (
        style.startswith("List Bullet")
        or style.startswith("List Number")
        or style == "List Paragraph"
    )


def _table_rows(table) -> list[list[str]]:
    return [[(cell.text or "").strip() for cell in row.cells] for row in table.rows]


def _render(doc) -> list[str]:
    """Walk the body in document order.

    python-docx's doc.paragraphs and doc.tables are SEPARATE collections, so
    reading both would append every table after all the prose, detached from
    the section it belonged to — scrambled output that still looks plausible.
    Iterating doc.element.body preserves the real order. Tags are namespaced,
    so match on the suffix.

    Two things are deliberately dropped without a trace: content inside a
    body-level w:sdt (a Word content control) is not a recognized tag here
    and is skipped, and a table nested inside a table cell is never
    descended into (_Cell.text joins only its own paragraphs). Both are
    rare enough, and lossy-by-nature enough, not to be worth the added
    complexity — but the omission is intentional, not an oversight.
    """
    lines: list[str] = []
    for child in doc.element.body:
        if child.tag.endswith("}p"):
            line = _paragraph_line(Paragraph(child, doc))
            if line:
                lines.append(line)
        elif child.tag.endswith("}tbl"):
            lines.extend(render_markdown_table(_table_rows(Table(child, doc))))
    return lines


class DocxExtractor:
    scopes = (DRIVE_READONLY,)
    services = ("drive",)

    def extract(self, services: dict, file_id: str, mime: str) -> ExtractedText:
        payload = execute(services["drive"].files().get(fileId=file_id, alt="media"))

        # A parse failure is a property of the file, not of the upstream
        # call, so it maps to SourceUnsupported (422) rather than
        # SourceUnavailable (502). Kept out of execute(), which exists to
        # normalize Google API failures.
        try:
            doc = Document(io.BytesIO(payload))
        except Exception as e:
            raise SourceUnsupported(f"docx could not be parsed: {type(e).__name__}") from e

        return ExtractedText(text="\n".join(_render(doc)))
