import io

from docx import Document

from src.sources.google_extractors.docx import DocxExtractor

DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class _FakeFiles:
    def __init__(self, payload):
        self._payload = payload

    def get(self, *, fileId, alt=None):
        return self

    def execute(self):
        return self._payload


class _FakeDrive:
    def __init__(self, payload):
        self._files = _FakeFiles(payload)

    def files(self):
        return self._files


def _bytes_of(doc) -> bytes:
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _extract(doc):
    payload = _bytes_of(doc)
    return DocxExtractor().extract({"drive": _FakeDrive(payload)}, "file123", DOCX_MIME_TYPE)


def test_heading_styles_become_markdown_levels():
    doc = Document()
    doc.add_paragraph("Big", style="Heading 1")
    doc.add_paragraph("Small", style="Heading 3")
    out = _extract(doc).text
    assert "# Big" in out
    assert "### Small" in out


def test_heading_7_clamps_to_the_deepest_markdown_level():
    doc = Document()
    doc.add_paragraph("Deep", style="Heading 7")
    out = _extract(doc).text
    assert "###### Deep" in out


def test_normal_paragraphs_have_no_prefix():
    doc = Document()
    doc.add_paragraph("just words")
    out = _extract(doc).text
    assert "just words" in out
    assert "#" not in out


def test_list_styled_paragraphs_get_a_dash():
    doc = Document()
    doc.add_paragraph("an item", style="List Bullet")
    assert "- an item" in _extract(doc).text


def test_list_bullet_2_and_list_number_get_a_dash():
    doc = Document()
    doc.add_paragraph("nested item", style="List Bullet 2")
    doc.add_paragraph("numbered item", style="List Number")
    out = _extract(doc).text
    assert "- nested item" in out
    assert "- numbered item" in out


def test_list_continue_and_plain_list_styles_are_not_bulleted():
    # "List Continue" and plain "List" contain the word "List" but are not
    # bullets: List Continue is prose continuing under a list item without
    # its own bullet, and List is just an indented paragraph.
    doc = Document()
    doc.add_paragraph("continuation prose", style="List Continue")
    doc.add_paragraph("indented text", style="List")
    out = _extract(doc).text
    assert "- continuation prose" not in out
    assert "continuation prose" in out
    assert "- indented text" not in out
    assert "indented text" in out


def test_a_table_renders_as_a_markdown_table():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Tier"
    table.cell(0, 1).text = "Amount"
    table.cell(1, 0).text = "Gold"
    table.cell(1, 1).text = "$5,000"
    out = _extract(doc).text
    assert "| Tier | Amount |" in out
    assert "| --- | --- |" in out
    assert "| Gold | $5,000 |" in out


def test_document_order_is_preserved_across_paragraphs_and_tables():
    # THE regression test for this extractor. python-docx exposes paragraphs
    # and tables as separate collections; reading both naively puts every
    # table after all the prose, silently.
    doc = Document()
    doc.add_paragraph("before the table")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "in the table"
    doc.add_paragraph("after the table")
    out = _extract(doc).text
    assert out.index("before the table") < out.index("in the table")
    assert out.index("in the table") < out.index("after the table")


def test_empty_paragraphs_do_not_emit_blank_lines():
    doc = Document()
    doc.add_paragraph("first")
    doc.add_paragraph("")
    doc.add_paragraph("second")
    assert "\n\n" not in _extract(doc).text


def test_extractor_declares_drive_scope_and_service():
    assert DocxExtractor().scopes == ("https://www.googleapis.com/auth/drive.readonly",)
    assert DocxExtractor().services == ("drive",)
